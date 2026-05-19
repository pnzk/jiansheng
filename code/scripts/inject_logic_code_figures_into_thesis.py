from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(r"e:\DATA\jianshen")
INPUT_DOCX = ROOT / "output" / "thesis_patch_output.docx"
OUTPUT_DOCX = ROOT / "output" / "基于大数据技术的健身房运动行为与健身效果分析-修改版_逻辑代码图版_v6.docx"
IMAGE_DIR = ROOT / "output" / "doc" / "images"


FIGURE_CONFIGS = [
    {
        "anchor": "图 5- 4  管理员监控与行为分析页面",
        "images": [
            {
                "path": IMAGE_DIR / "code_fig_5_4a_admin_monitor_metrics.png",
                "label": "（a）监控指标与活跃率计算",
            },
            {
                "path": IMAGE_DIR / "code_fig_5_4b_admin_behavior_retention.png",
                "label": "（b）行为趋势与留存率计算",
            },
        ],
        "caption": "图5-4 管理员监控与行为分析页面逻辑实现代码截图",
        "description": "图5-4中，子图（a）展示了管理员监控页面所依赖的核心统计逻辑，主要完成活跃用户数、平均活跃率和相关监控指标的统一计算，用于支撑全局看板中的关键数值卡片。子图（b）展示了用户行为分析逻辑的核心实现，系统在该过程中构造日级活动明细、行为趋势数据以及留存率结果，并将其统一封装后返回前端。因此，管理员端的监控与行为分析能力并非简单页面拼装，而是建立在统一统计服务和稳定数据回退机制基础上的综合分析实现。",
    },
    {
        "anchor": "图 5- 5学员管理页面",
        "images": [
            {
                "path": IMAGE_DIR / "code_fig_5_5_admin_students.png",
                "label": "",
            }
        ],
        "caption": "图5-5 学员管理页面逻辑实现代码截图",
        "description": "图5-5展示了学员管理页面的核心业务逻辑。系统首先按照学生角色筛选用户，再提取教练编号并建立教练姓名映射关系，最后将账号、姓名、联系方式、健身目标、教练归属等字段统一封装为 StudentResponse 返回前端页面。由此可见，学员管理页面所展示的数据并不是零散表字段的直接拼接，而是经过服务层统一组织后的结果，这样既保证了主数据展示的一致性，也保证了管理员在进行学员管理时能够获得稳定完整的数据支撑。",
    },
    {
        "anchor": "图 5- 6  教练总览与学员报告页面",
        "images": [
            {
                "path": IMAGE_DIR / "code_fig_5_6a_coach_dashboard.png",
                "label": "（a）教练总览统计逻辑",
            },
            {
                "path": IMAGE_DIR / "code_fig_5_6b_coach_student_report.png",
                "label": "（b）学员报告聚合逻辑",
            },
        ],
        "caption": "图5-6 教练总览与学员报告页面逻辑实现代码截图",
        "description": "图5-6中，子图（a）给出了教练总览页面的关键统计逻辑，系统围绕所属学员总数、性别结构、平均年龄、目标分布和活跃学员数量等指标构建教练工作看板，用于帮助教练快速掌握整体学员状态。子图（b）则展示了学员报告的聚合逻辑，系统进一步联合体测记录、运动记录和训练计划进度，生成体重变化、训练时长、热量消耗和计划完成度等分析结果。这说明教练端不仅具备群体概览能力，也具备面向具体学员输出分析报告的实现基础。",
    },
]


def insert_paragraph_after(paragraph, text="", align=WD_ALIGN_PARAGRAPH.CENTER):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def clear_paragraph(paragraph):
    paragraph.clear()


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False):
    run.bold = bold
    run.font.name = latin
    run.font.size = int(size_pt * 12700)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def apply_paragraph_font(paragraph, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False):
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, latin=latin, size_pt=size_pt, bold=bold)


def replace_placeholder_block(doc: Document, config: dict):
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == config["anchor"]:
            placeholder_caption = paragraphs[index + 2] if index + 2 < len(paragraphs) else None
            placeholder_desc = paragraphs[index + 3] if index + 3 < len(paragraphs) else None

            current = paragraph
            for image_item in config["images"]:
                pic_para = insert_paragraph_after(current, align=WD_ALIGN_PARAGRAPH.CENTER)
                pic_para.add_run().add_picture(str(image_item["path"]), width=Cm(15.2))
                current = pic_para
                if image_item.get("label"):
                    label_para = insert_paragraph_after(current, image_item["label"], align=WD_ALIGN_PARAGRAPH.CENTER)
                    apply_paragraph_font(label_para, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False)
                    current = label_para

            if placeholder_caption is not None:
                clear_paragraph(placeholder_caption)
                placeholder_caption.add_run(config["caption"])
                placeholder_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_paragraph_font(placeholder_caption, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False)

            if placeholder_desc is not None:
                clear_paragraph(placeholder_desc)
                placeholder_desc.add_run(config["description"])
                placeholder_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                apply_paragraph_font(placeholder_desc, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False)

            return True
    return False


def main():
    doc = Document(str(INPUT_DOCX))
    replaced = []
    for config in FIGURE_CONFIGS:
        if replace_placeholder_block(doc, config):
            replaced.append(config["caption"])

    doc.save(str(OUTPUT_DOCX))
    print(f"[OK] saved {OUTPUT_DOCX}")
    print(f"[OK] replaced {len(replaced)} figure blocks")
    for item in replaced:
        print(f" - {item}")


if __name__ == "__main__":
    main()
