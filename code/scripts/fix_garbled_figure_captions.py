from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"e:\DATA\jianshen")
INPUT_DOCX = ROOT / "output" / "基于大数据技术的健身房运动行为与健身效果分析-修改版_逻辑代码图版_v6.docx"
OUTPUT_DOCX = ROOT / "output" / "基于大数据技术的健身房运动行为与健身效果分析-修改版_逻辑代码图版_v7.docx"


REPLACEMENTS = {
    299: "图5-1 数据来源网站与接口配置核心代码截图",
    300: "图5-1反映了系统在实现层面对“数据来源多样化”的具体落地方式。通过将来源站点、接口入口与配置参数集中维护，后续的数据拉取、失败回退和样本扩展都具备了明确依托，这为健身场景下多源数据融合提供了统一起点。",
    306: "图5-2 运动行为样本清洗核心代码截图",
    307: "图5-2所示代码展示了运动行为样本清洗的核心实现，包括重复记录剔除、活动日期解析、活动分钟阈值判定、运动类型归类以及心率与热量估算等步骤。该实现直接对应第四章中有效运动样本判定与能耗估算的设计思想，是后续运动记录统计与分析结果可靠性的基础。",
    313: "图5-3 体测数据清洗与指标补齐核心代码截图",
    314: "图5-3展示了体测数据清洗与指标补齐的关键实现过程，系统在该部分完成日期解析、体重值有效性判定、身高补齐、BMI 自动计算以及肌肉量估算等处理，从而保证体测样本能够以统一口径进入分析流程，并为后续健身效果评估提供可持续使用的数据基础。",
}


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size_pt=10.5, bold=False):
    run.bold = bold
    run.font.name = latin
    run.font.size = int(size_pt * 12700)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def apply_font(paragraph, size_pt=10.5, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt, bold=bold)


def main():
    doc = Document(str(INPUT_DOCX))
    for index, new_text in REPLACEMENTS.items():
        paragraph = doc.paragraphs[index]
        paragraph.clear()
        paragraph.add_run(new_text)
        if index in {299, 306, 313}:
            paragraph.alignment = 1
            apply_font(paragraph, size_pt=10.5, bold=False)
        else:
            paragraph.alignment = 3
            apply_font(paragraph, size_pt=10.5, bold=False)

    doc.save(str(OUTPUT_DOCX))
    print(f"[OK] saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
